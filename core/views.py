from django.shortcuts import render,redirect, get_object_or_404
from .forms import RegisterForm, UpdateProfileForm,PaymentForm,ConfirmPaymentForm
from django.contrib import messages
from django.contrib.auth import authenticate, login,logout
from django.contrib.auth.decorators import login_required
from .models import StationPrepay, Communication, Payment


def signup(request):
    form = RegisterForm()
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Account created successfully")
            return redirect("signup")
        
    context = {"form":form}
    return render(request, "core/signup.html", context)

def signin (request):
    if request.method == 'POST':
        email = request.POST["email"]
        password= request.POST["password"]

        user = authenticate(request, email=email, password=password)
        if user is not None:
            login(request, user)
            return redirect('index')
    context= {}
    return render(request, "core/login.html", context)

def signout(request):
    logout(request)
    return redirect("index")

@login_required(login_url="signin")
def profile(request):
    user = request.user
    
    # Récupérer les paiements validés associés à l'utilisateur
    payments = Payment2.objects.filter(user=user, payement_valide=True).order_by('-created_at')

    context = {
        "user": user,
        "payments": payments,  # Passer les paiements à la vue
    }
    return render(request, "core/profile.html", context)

@login_required(login_url="signin")
def update_profile(request):
    if request.user.is_authenticated:
        user = request.user
        form = UpdateProfileForm(instance=user)
        if request.method == 'POST':
            form = UpdateProfileForm(request.POST, request.FILES, instance=user)
            if form.is_valid():
                form.save()
                messages.success(request, "Profile updated successfully")
                return redirect("profile")
                
    context = {"form": form}
    return render(request, "core/update_profile.html", context)

from django.shortcuts import render, redirect
from django.contrib.admin.views.decorators import staff_member_required
from .models import Payment, PaymentConfirmation
from .forms import PaymentConfirmationForm

@staff_member_required
def confirm_payment(request, payment_id):
    # Récupérer le paiement à confirmer
    payment = get_object_or_404(Payment, id=payment_id)

    # Si l'utilisateur n'est pas admin, rediriger
    if not request.user.is_superuser:
        messages.error(request, "Vous n'avez pas les permissions nécessaires.")
        return redirect('index')

    # Si le formulaire est soumis
    if request.method == 'POST':
        form = ConfirmPaymentForm(request.POST)
        if form.is_valid():
            # Récupérer l'ID de la transaction et le numéro opérant
            transaction_id = form.cleaned_data['transaction_id']
            operant_number = form.cleaned_data['operant_number']
            
            # Marquer le paiement comme confirmé
            payment.transaction_id = transaction_id
            payment.operant_number = operant_number
            payment.status = Payment.STATUS_CONFIRMED  # Assure-toi que ce statut est défini dans le modèle
            payment.save()

            # Ajouter un message de succès et rediriger
            messages.success(request, "Le paiement a été confirmé avec succès.")
            return redirect('index')
    else:
        form = ConfirmPaymentForm()

    context = {
        'payment': payment,
        'form': form,
    }
    return render(request, 'core/confirm_payment.html', context)

from django.shortcuts import render, get_object_or_404, redirect
from .models import StationPrepay, Communication, Cart
from django.contrib.auth.decorators import login_required
from django.db.models.functions import TruncDate

@login_required
def station_detail(request, slug):
    # Récupérer la station par son slug
    station = get_object_or_404(StationPrepay, slug=slug)

    # Récupérer les communications associées à la station
    communications = Communication.objects.filter(station=station)

    # Récupérer les numéros de paiement associés à la station
    payment_numbers = station.payment_numbers.all()

     # Récupérer les paiements par date
    payments_by_date = Payment2.objects.filter(cart__station_prepay=station) \
                                        .annotate(payment_date=TruncDate('created_at')) \
                                        .values('payment_date') \
                                        .distinct() 

    # Si le formulaire du panier est soumis
    if request.method == 'POST':
        montant = request.POST.get('montant')  # Montant du panier

        try:
            montant = float(montant)
        except ValueError:
            return render(request, 'core/station_detail.html', {
                'station': station,
                'communications': communications,
                'payment_numbers': payment_numbers,
                'error_message': "Le montant est invalide, veuillez réessayer."
            })

        # Créer un panier pour l'utilisateur et la station
        cart = Cart.objects.create(
            station_prepay=station,
            montant=montant,
            user=request.user
        )

        # Rediriger vers la page de checkout
        return redirect('checkout', cart_id=cart.id, slug=station.slug)

    context = {
        'station': station,
        'communications': communications,
        'payment_numbers': payment_numbers,
        'payments_by_date': payments_by_date
    }

    return render(request, 'core/station_detail.html', context)

from django.http import HttpResponse
from docx import Document
from django.shortcuts import render, get_object_or_404
from .models import StationPrepay, Payment2
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils.dateparse import parse_date

@login_required
def payments_by_date(request, slug, payment_date=None):
    # Récupérer la station par son slug
    station = get_object_or_404(StationPrepay, slug=slug)

    # Filtrer les paiements validés pour la station
    payments = Payment2.objects.filter(cart__station_prepay=station, payement_valide=True)

    # Si une date spécifique est donnée, filtrer les paiements par date
    if payment_date:
        try:
            # Convertir la date au format attendu (année-mois-jour)
            date_object = parse_date(payment_date)
            payments = payments.filter(created_at__date=date_object)
        except ValueError:
            messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")

    # Trier les paiements par date (du plus récent au plus ancien)
    payments = payments.order_by('created_at')

    # Calculer le nombre de paiements et le montant total
    payment_count = payments.count()  # Nombre de paiements
    total_amount = sum(payment.montant for payment in payments)  # Montant total des paiements

    # Ajouter une option pour générer un fichier Word
    if 'generate_word' in request.GET:
        return generate_word(request, station, payment_date, payment_count, total_amount, payments)

    context = {
        'station': station,
        'payments': payments,
        'payment_date': payment_date,
        'payment_count': payment_count,
        'total_amount': total_amount,
    }

    return render(request, 'core/payments_by_date.html', context)


def generate_word(request, station, payment_date, payment_count, total_amount, payments):
    # Créer un document Word
    doc = Document()

    # Ajouter un titre
    doc.add_heading(f'Paiements pour le {payment_date} - {station.nom}', 0)

    # Ajouter le nombre de paiements et le montant total
    doc.add_paragraph(f'Nombre de paiements : {payment_count}')
    doc.add_paragraph(f'Montant total des paiements : {total_amount} FC')

    # Ajouter un tableau pour afficher les paiements
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nom du client'
    hdr_cells[1].text = 'Montant'
    hdr_cells[2].text = 'Numéro de transfert'
    hdr_cells[3].text = 'Identifiant client'
    hdr_cells[4].text = 'Statut'

    # Ajouter les paiements dans le tableau
    for payment in payments:
        row_cells = table.add_row().cells
        row_cells[0].text = f'{payment.nom} {payment.prenom}'
        row_cells[1].text = str(payment.montant)
        row_cells[2].text = str(payment.numero_transfert)
        row_cells[3].text = str(payment.numero_identifiant)
        row_cells[4].text = 'Validé' if payment.payement_valide else 'Non Validé'

    # Créer une réponse HTTP avec le fichier Word en tant que contenu
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="paiements_{payment_date}_{station.nom}.docx"'

    # Sauver le document dans la réponse HTTP
    doc.save(response)

    return response


# def station_detail(request, slug):
#     # Récupérer la station par son slug
#     station = get_object_or_404(StationPrepay, slug=slug)
    
#     # Récupérer les communications associées à la station
#     communications = Communication.objects.filter(station=station)
    
#     # Récupérer les numéros de paiement associés à la station
#     payment_numbers = station.payment_numbers.all()  # Tu sembles avoir une relation avec les numéros de paiement
    
#     # Si l'utilisateur soumet un paiement
#        # Passer les objets à la template
#     context = {
#         'station': station,
#         'communications': communications,
        
#         'payment_numbers': payment_numbers,  # On passe aussi les numéros de paiement dans le contexte
#     }
    
#     return render(request, 'core/station_detail.html', context) # Le formulaire que nous allons créer pour gérer les infos
# views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Cart, Payment2
from .forms import CheckoutForm
import logging

logger = logging.getLogger(__name__)

@login_required
def checkout(request, cart_id,slug):
    # Récupérer le panier associé à l'utilisateur
    cart = get_object_or_404(Cart, id=cart_id, user=request.user)
    # Récupérer la station par son slug
    station = get_object_or_404(StationPrepay, slug=slug)
    payment_numbers = station.payment_numbers.all()

    logger.debug(f"Montant: {cart.montant}, Commission: {cart.commission}, Total: {cart.total}")

    if request.method == 'POST':
        form = CheckoutForm(request.POST)
        if form.is_valid():
            # Sauvegarder le paiement en lien avec le panier et l'utilisateur
            payment = form.save(commit=False)
            payment.cart = cart  # Lier le paiement au panier
            payment.user = request.user  # Lier le paiement à l'utilisateur
            payment.save()  # Sauvegarder le paiement

            # Afficher un message de succès
            messages.success(request, 'Votre paiement a bien été enregistré. Merci !')

            # Rediriger vers une page de confirmation
            return redirect('payment_success', payment_id=payment.id)
        else:
            # Si le formulaire n'est pas valide, afficher un message d'erreur
            messages.error(request, 'Veuillez vérifier les informations soumises.')

    else:
        # Initialiser le formulaire vide
        form = CheckoutForm()

    return render(request, 'core/checkout.html', {'form': form, 'cart': cart,'payment_numbers': payment_numbers})



def payment_success(request, payment_id):
    # Récupérer l'objet Payment2
    payment = get_object_or_404(Payment2, id=payment_id)

    # Afficher la page de succès avec les détails du paiement
    return render(request, 'core/payment_success.html', {'payment': payment})

from django.shortcuts import render, get_object_or_404
from django.db.models import DateField
from django.db.models.functions import TruncDate
from .models import StationPrepay, Payment2
from django.utils.dateparse import parse_date
from django.contrib import messages
from .models import Payment2, StationPrepay

@login_required
def station_payments(request, slug):
    # Récupérer la station par son slug
    station = get_object_or_404(StationPrepay, slug=slug)
    
    # Récupérer les paiements par date
    payments_by_date = Payment2.objects.filter(cart__station_prepay=station) \
                                        .annotate(payment_date=TruncDate('created_at')) \
                                        .values('payment_date') \
                                        .distinct()

    # Si la barre de recherche est utilisée
    search_date = request.GET.get('search_date')
    if search_date:
        # Convertir la date en format valide
        try:
            search_date = parse_date(search_date)
            payments_by_date = payments_by_date.filter(payment_date=search_date)
        except ValueError:
            # Gérer le cas où la date est invalide
            messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
            payments_by_date = []

    context = {
        'station': station,
        'payments_by_date': payments_by_date,
        'search_date': search_date,
    }

    return render(request, 'core/station_payments.html', context)

from django.http import HttpResponse
from docx import Document
from django.shortcuts import render, get_object_or_404
from .models import StationPrepay, Payment2
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils.dateparse import parse_date

@login_required
def payments_by_date(request, slug, payment_date=None):
    # Récupérer la station par son slug
    station = get_object_or_404(StationPrepay, slug=slug)

    # Filtrer les paiements validés pour la station
    payments = Payment2.objects.filter(cart__station_prepay=station, payement_valide=True)

    # Si une date spécifique est donnée, filtrer les paiements par date
    if payment_date:
        try:
            # Convertir la date au format attendu (année-mois-jour)
            date_object = parse_date(payment_date)
            payments = payments.filter(created_at__date=date_object)
        except ValueError:
            messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")

    # Trier les paiements par date (du plus récent au plus ancien)
    payments = payments.order_by('created_at')

    # Calculer le nombre de paiements et le montant total
    payment_count = payments.count()  # Nombre de paiements
    total_amount = sum(payment.montant for payment in payments)  # Montant total des paiements

    # Si la requête contient le paramètre 'generate_word', générer et télécharger le fichier Word
    if 'generate_word' in request.GET:
        return generate_word(request, station, payment_date, payment_count, total_amount, payments)

    context = {
        'station': station,
        'payments': payments,
        'payment_date': payment_date,
        'payment_count': payment_count,
        'total_amount': total_amount,
    }

    return render(request, 'core/payments_by_date.html', context)


def generate_word(request, station, payment_date, payment_count, total_amount, payments):
    # Créer un document Word
    doc = Document()

    # Ajouter un titre
    doc.add_heading(f'Rapport des paiements pour la station: {station.nom}', 0)

    # Ajouter la date du rapport
    if payment_date:
        doc.add_paragraph(f'Date du rapport: {payment_date}')
    else:
        doc.add_paragraph('Date du rapport: Tous les paiements')

    # Ajouter un résumé des paiements
    doc.add_paragraph(f'Nombre total de paiements : {payment_count}')
    doc.add_paragraph(f'Montant total des paiements : {total_amount} FC')

    # Ajouter un tableau des paiements
    table = doc.add_table(rows=1, cols=5)
    headers = ['Nom du Client', 'Numéro de Transfert', 'Montant', 'Identifiant Client', 'Statut']

    # Ajouter les en-têtes de colonnes
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Ajouter les paiements au tableau
    for payment in payments:
        row = table.add_row().cells
        row[0].text = f'{payment.nom} {payment.prenom}'
        row[1].text = str(payment.numero_transfert)
        row[2].text = f'{payment.montant} FC'
        row[3].text = str(payment.numero_identifiant)
        row[4].text = 'Validé' if payment.payement_valide else 'Non Validé'

    # Créer la réponse HTTP pour télécharger le fichier Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{station.nom}_paiements_{payment_date or "all"}_rapport.docx"'

    # Sauvegarder le document dans la réponse
    doc.save(response)

    return response

from .models import BaseClientStation

def verifier_identifiant(request, slug):
    # Récupérer la station par son slug
    station = get_object_or_404(StationPrepay, slug=slug)
    
    client = None  # Variable pour stocker le client si trouvé
    if request.method == "POST":
        identifiant = request.POST.get('identifiant')

        try:
            # Recherche de l'identifiant dans la base de données pour cette station spécifique
            client = BaseClientStation.objects.get(identifiant=identifiant, station_prepay=station)
            messages.success(request, f"Identifiant trouvé pour {client.nom} {client.prenom}.")
        except BaseClientStation.DoesNotExist:
            messages.error(request, "Votre identifiant n'existe pas chez nous pour cette station. Veuillez vérifier.")

    return render(request, 'core/verifier_identifiant.html', {'station': station, 'client': client})